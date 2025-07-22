#!/usr/bin/env python3
"""
File Generator Agent
Genera archivos finales (Excel, Word) con los resultados consolidados
"""
import os
import json
from typing import Dict, List, Optional, Any
from pathlib import Path
from datetime import datetime
import tkinter as tk
from tkinter import filedialog, messagebox
from safe_print_utils import safe_print_global as safe_print

try:
    import openpyxl
    from openpyxl.styles import Font, Alignment, PatternFill
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False
    safe_print("[WARNING] openpyxl no disponible - instalando...")

try:
    from docx import Document
    from docx.shared import Inches
    WORD_AVAILABLE = True
except ImportError:
    WORD_AVAILABLE = False
    safe_print("[WARNING] python-docx no disponible - instalando...")

class FileGenerator:
    """
    Genera archivos finales con los resultados consolidados
    """
    
    def __init__(self):
        self.install_dependencies()
    
    def install_dependencies(self):
        """
        Instala las dependencias necesarias si no están disponibles
        """
        global EXCEL_AVAILABLE, WORD_AVAILABLE
        
        if not EXCEL_AVAILABLE:
            try:
                import subprocess
                subprocess.check_call(['pip', 'install', 'openpyxl'])
                import openpyxl
                from openpyxl.styles import Font, Alignment, PatternFill
                EXCEL_AVAILABLE = True
                safe_print("[SUCCESS] openpyxl instalado correctamente")
            except Exception as e:
                safe_print(f"[ERROR] No se pudo instalar openpyxl: {e}")
        
        if not WORD_AVAILABLE:
            try:
                import subprocess
                subprocess.check_call(['pip', 'install', 'python-docx'])
                from docx import Document
                WORD_AVAILABLE = True
                safe_print("[SUCCESS] python-docx instalado correctamente")
            except Exception as e:
                safe_print(f"[ERROR] No se pudo instalar python-docx: {e}")
    
    def choose_output_file(self, default_name: str = "resultado", file_type: str = "excel") -> Optional[str]:
        """
        Muestra diálogo para que el usuario elija el destino del archivo
        """
        try:
            # Create hidden root window
            root = tk.Tk()
            root.withdraw()  # Hide main window
            root.lift()
            root.attributes('-topmost', True)
            
            # Configure file types based on requested type
            if file_type.lower() == "excel":
                filetypes = [("Excel files", "*.xlsx"), ("All files", "*.*")]
                default_ext = ".xlsx"
            elif file_type.lower() == "word":
                filetypes = [("Word documents", "*.docx"), ("All files", "*.*")]
                default_ext = ".docx"
            else:
                filetypes = [("All files", "*.*")]
                default_ext = ".txt"
            
            # Generate default name with timestamp
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            default_filename = f"{default_name}_{timestamp}{default_ext}"
            
            # Show save file dialog
            file_path = filedialog.asksaveasfilename(
                title=f"Save {file_type} file",
                defaultextension=default_ext,
                filetypes=filetypes,
                initialfile=default_filename
            )
            
            root.destroy()
            
            if file_path:
                safe_print(f"[USER] File selected: {file_path}")
                return file_path
            else:
                safe_print("[INFO] User cancelled file selection")
                return None
                
        except Exception as e:
            safe_print(f"[ERROR] Error showing file dialog: {e}")
            return None
    
    def generate_excel_file(self, consolidated_data: str, original_objective: str, 
                          summary_info: Dict[str, Any], output_path: str = None) -> Optional[str]:
        """
        Generate Excel file with consolidated results
        """
        if not EXCEL_AVAILABLE:
            safe_print("[ERROR] openpyxl not available to generate Excel")
            return None
        
        try:
            # Seleccionar archivo de salida si no se proporciona
            if not output_path:
                output_path = self.choose_output_file("extracted_results", "excel")
                if not output_path:
                    return None
            
            safe_print(f"[EXCEL] Generando archivo Excel: {output_path}")
            
            # Crear workbook
            from openpyxl import Workbook
            from openpyxl.styles import Font, Alignment, PatternFill
            
            wb = Workbook()
            
            # Summary sheet
            ws_summary = wb.active
            ws_summary.title = "Summary"
            
            # Summary header
            header_fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
            header_font = Font(color="FFFFFF", bold=True, size=14)
            
            ws_summary['A1'] = "WEB DATA EXTRACTION REPORT"
            ws_summary['A1'].font = header_font
            ws_summary['A1'].fill = header_fill
            ws_summary.merge_cells('A1:D1')
            
            # General information
            row = 3
            info_font = Font(bold=True)
            
            ws_summary[f'A{row}'] = "Original Objective:"
            ws_summary[f'A{row}'].font = info_font
            ws_summary[f'B{row}'] = original_objective
            row += 1
            
            ws_summary[f'A{row}'] = "Extraction Date:"
            ws_summary[f'A{row}'].font = info_font
            ws_summary[f'B{row}'] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            row += 1
            
            ws_summary[f'A{row}'] = "Pages Processed:"
            ws_summary[f'A{row}'].font = info_font
            ws_summary[f'B{row}'] = summary_info.get('pages_processed', 0)
            row += 1
            
            ws_summary[f'A{row}'] = "Total Characters:"
            ws_summary[f'A{row}'].font = info_font
            ws_summary[f'B{row}'] = summary_info.get('total_content_chars', 0)
            row += 2
            
            # Consolidated results
            ws_summary[f'A{row}'] = "CONSOLIDATED RESULTS:"
            ws_summary[f'A{row}'].font = info_font
            row += 1
            
            # Split consolidated text into lines and add them
            for line in consolidated_data.split('\n'):
                if line.strip():
                    ws_summary[f'A{row}'] = line.strip()
                    row += 1
            
            # Page details sheet
            if summary_info.get('pages_info'):
                ws_details = wb.create_sheet("Page Details")
                
                # Headers
                headers = ['Page', 'Title', 'URL', 'Characters']
                for col, header in enumerate(headers, 1):
                    cell = ws_details.cell(row=1, column=col, value=header)
                    cell.font = header_font
                    cell.fill = header_fill
                
                # Page data
                for row, page_info in enumerate(summary_info['pages_info'], 2):
                    ws_details.cell(row=row, column=1, value=page_info.get('page_number', 'N/A'))
                    ws_details.cell(row=row, column=2, value=page_info.get('title', 'N/A'))
                    ws_details.cell(row=row, column=3, value=page_info.get('url', 'N/A'))
                    ws_details.cell(row=row, column=4, value=page_info.get('content_length', 0))
                
                # Ajustar ancho de columnas
                try:
                    for col_idx in range(1, 5):  # Columnas A-D
                        max_length = 0
                        column_letter = chr(64 + col_idx)  # A=1, B=2, etc.
                        
                        for row_idx in range(1, ws_details.max_row + 1):
                            cell = ws_details.cell(row=row_idx, column=col_idx)
                            try:
                                if cell.value and len(str(cell.value)) > max_length:
                                    max_length = len(str(cell.value))
                            except:
                                pass
                                
                        adjusted_width = min(max_length + 2, 50)
                        ws_details.column_dimensions[column_letter].width = adjusted_width
                except Exception as e:
                    safe_print(f"[WARNING] Error ajustando ancho de columnas: {e}")
            
            # Ajustar ancho de columnas en resumen
            try:
                for col_idx in range(1, 5):  # Columnas A-D
                    max_length = 0
                    column_letter = chr(64 + col_idx)  # A=1, B=2, etc.
                    
                    for row_idx in range(1, ws_summary.max_row + 1):
                        cell = ws_summary.cell(row=row_idx, column=col_idx)
                        try:
                            if cell.value and len(str(cell.value)) > max_length:
                                max_length = len(str(cell.value))
                        except:
                            pass
                            
                    adjusted_width = min(max_length + 2, 80)
                    ws_summary.column_dimensions[column_letter].width = adjusted_width
            except Exception as e:
                safe_print(f"[WARNING] Error ajustando ancho de columnas en resumen: {e}")
            
            # Guardar archivo
            wb.save(output_path)
            
            safe_print(f"[SUCCESS] Archivo Excel generado: {output_path}")
            return output_path
            
        except Exception as e:
            safe_print(f"[ERROR] Error generando archivo Excel: {e}")
            return None
    
    def generate_word_file(self, consolidated_data: str, original_objective: str,
                         summary_info: Dict[str, Any], output_path: str = None) -> Optional[str]:
        """
        Generate Word file with consolidated results
        """
        if not WORD_AVAILABLE:
            safe_print("[ERROR] python-docx not available to generate Word")
            return None
        
        try:
            # Select output file if not provided
            if not output_path:
                output_path = self.choose_output_file("extracted_report", "word")
                if not output_path:
                    return None
            
            safe_print(f"[WORD] Generating Word file: {output_path}")
            
            # Create document
            from docx import Document
            from docx.shared import Inches
            
            doc = Document()
            
            # Main title
            title = doc.add_heading('WEB DATA EXTRACTION REPORT', 0)
            title.alignment = 1  # Centered
            
            # General information
            doc.add_heading('General Information', level=1)
            
            info_para = doc.add_paragraph()
            info_para.add_run('Original Objective: ').bold = True
            info_para.add_run(original_objective)
            
            info_para = doc.add_paragraph()
            info_para.add_run('Extraction Date: ').bold = True
            info_para.add_run(datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
            
            info_para = doc.add_paragraph()
            info_para.add_run('Pages Processed: ').bold = True
            info_para.add_run(str(summary_info.get('pages_processed', 0)))
            
            info_para = doc.add_paragraph()
            info_para.add_run('Total Characters Extracted: ').bold = True
            info_para.add_run(f"{summary_info.get('total_content_chars', 0):,}")
            
            # Consolidated results
            doc.add_heading('Consolidated Results', level=1)
            
            # Agregar el contenido consolidado
            for line in consolidated_data.split('\n'):
                if line.strip():
                    doc.add_paragraph(line.strip())
            
            # Detalles por página (si hay información)
            if summary_info.get('pages_info'):
                doc.add_heading('Detalles por Página', level=1)
                
                # Crear tabla
                table = doc.add_table(rows=1, cols=4)
                table.style = 'Light Grid Accent 1'
                
                # Encabezados
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Página'
                hdr_cells[1].text = 'Título'
                hdr_cells[2].text = 'URL'
                hdr_cells[3].text = 'Caracteres'
                
                # Datos
                for page_info in summary_info['pages_info']:
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(page_info.get('page_number', 'N/A'))
                    row_cells[1].text = page_info.get('title', 'N/A')
                    row_cells[2].text = page_info.get('url', 'N/A')
                    row_cells[3].text = str(page_info.get('content_length', 0))
            
            # Pie de página
            doc.add_paragraph('\n\n')
            footer_para = doc.add_paragraph()
            footer_para.add_run('Generado automáticamente por Web Agent - ').italic = True
            footer_para.add_run(datetime.now().strftime("%Y-%m-%d %H:%M:%S")).italic = True
            footer_para.alignment = 1  # Centrado
            
            # Guardar archivo
            doc.save(output_path)
            
            safe_print(f"[SUCCESS] Archivo Word generado: {output_path}")
            return output_path
            
        except Exception as e:
            safe_print(f"[ERROR] Error generando archivo Word: {e}")
            return None
    
    def open_file_location(self, file_path: str):
        """
        Abre la ubicación del archivo generado
        """
        try:
            import subprocess
            import platform
            
            if platform.system() == 'Windows':
                subprocess.run(['explorer', '/select,', file_path])
            elif platform.system() == 'Darwin':  # macOS
                subprocess.run(['open', '-R', file_path])
            else:  # Linux
                subprocess.run(['xdg-open', os.path.dirname(file_path)])
                
            safe_print(f"[INFO] File location opened: {file_path}")
            
        except Exception as e:
            safe_print(f"[ERROR] Could not open file location: {e}")
    
    def show_success_dialog(self, file_path: str):
        """
        Show success dialog with option to open the file
        """
        try:
            root = tk.Tk()
            root.withdraw()
            
            result = messagebox.askyesno(
                "File Generated Successfully",
                f"The file has been generated correctly:\n\n{file_path}\n\nDo you want to open the file location?",
                icon="info"
            )
            
            root.destroy()
            
            if result:
                self.open_file_location(file_path)
                
        except Exception as e:
            safe_print(f"[ERROR] Error showing success dialog: {e}")
